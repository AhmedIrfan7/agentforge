"""Unit tests for extraction_metadata.py (roadmap step 094) --
build_doc_metadata's cleaning/merging logic and detect_language's
confidence threshold, exercised directly against real py3langid (a real,
if small, ML model -- not mocked, same "no mocks for real infrastructure"
stance this project takes for external services)."""

import io

import docx

from extraction_metadata import build_doc_metadata, detect_language, normalize_language_code


def test_detects_confident_english_text() -> None:
    text = "This is a perfectly ordinary English sentence about quarterly revenue and growth."
    assert detect_language(text) == "en"


def test_detects_confident_french_text() -> None:
    text = (
        "Ceci est une phrase tout a fait ordinaire en francais sur "
        "les revenus trimestriels et la croissance."
    )
    assert detect_language(text) == "fr"


def test_empty_text_returns_none_not_a_guess() -> None:
    assert detect_language("") is None
    assert detect_language("   ") is None


def test_low_confidence_garbage_returns_none_not_a_guess() -> None:
    # py3langid's norm_probs=True classifier converges to the same low
    # ~0.17 confidence for degenerate input (verified live) -- below
    # the threshold, this must not report a language it isn't sure of.
    assert detect_language("123456 !@#$%^&*()") is None


def test_normalize_language_code_strips_region_subtag() -> None:
    assert normalize_language_code("en-US") == "en"
    assert normalize_language_code("en_US") == "en"
    assert normalize_language_code("FR") == "fr"


def test_build_doc_metadata_for_unregistered_extension_only_detects_language() -> None:
    text = "This is real English content for a plain text file."
    metadata = build_doc_metadata("txt", text.encode("utf-8"), text)
    assert metadata["title"] is None
    assert metadata["author"] is None
    assert metadata["created_at"] is None
    assert metadata["modified_at"] is None
    assert metadata["language"] == "en"


def test_build_doc_metadata_filters_known_author_sentinels() -> None:
    # A brand-new python-docx Document(), never given a real author,
    # defaults core_properties.author to the literal string
    # "python-docx" -- verified live. Reporting that as "the author"
    # would be worse than reporting none at all.
    document = docx.Document()
    document.add_paragraph("Some content.")
    buf = io.BytesIO()
    document.save(buf)
    content = buf.getvalue()

    metadata = build_doc_metadata("docx", content, "Some content.")
    assert metadata["author"] is None


def test_build_doc_metadata_declared_language_wins_over_detection() -> None:
    # A document explicitly declaring French but written in English
    # (an artificial mismatch to make the point) should report the
    # declared language, not the detected one -- an author's own
    # declaration is trusted over a statistical guess.
    document = docx.Document()
    document.core_properties.language = "fr-FR"
    document.add_paragraph("This paragraph is actually written in English.")
    buf = io.BytesIO()
    document.save(buf)
    content = buf.getvalue()

    metadata = build_doc_metadata("docx", content, "This paragraph is actually written in English.")
    assert metadata["language"] == "fr"
