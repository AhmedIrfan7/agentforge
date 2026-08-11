"""Unit tests for extraction_docx.py (roadmap step 092) -- real .docx
files built with python-docx itself (it can write as well as read, so
no separate fixture-generation dependency is needed the way reportlab
was for PDF).
"""

import io
from collections.abc import Callable

import docx
from docx.document import Document

from extraction_docx import extract_docx


def _build_docx(build: Callable[[Document], None]) -> bytes:
    document = docx.Document()
    build(document)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_plain_paragraph_extracts_as_plain_text() -> None:
    def build(d: Document) -> None:
        d.add_paragraph("Just an ordinary sentence.")

    text = extract_docx(_build_docx(build))
    assert text == "Just an ordinary sentence."


def test_title_style_becomes_h1() -> None:
    def build(d: Document) -> None:
        d.add_heading("Document Title", level=0)
        d.add_paragraph("Body underneath.")

    text = extract_docx(_build_docx(build))
    assert "# Document Title" in text
    assert "## Document Title" not in text


def test_heading_levels_map_to_matching_hash_counts() -> None:
    def build(d: Document) -> None:
        d.add_heading("Top Section", level=1)
        d.add_heading("Sub Section", level=2)
        d.add_heading("Sub Sub Section", level=3)

    lines = extract_docx(_build_docx(build)).splitlines()
    assert "# Top Section" in lines
    assert "## Sub Section" in lines
    assert "### Sub Sub Section" in lines


def test_table_becomes_markdown_pipe_table_without_duplication() -> None:
    def build(d: Document) -> None:
        d.add_heading("People", level=2)
        table = d.add_table(rows=3, cols=2)
        for row, (name, age) in zip(
            table.rows, [("Name", "Age"), ("Alice", "30"), ("Bob", "25")], strict=True
        ):
            row.cells[0].text = name
            row.cells[1].text = age
        d.add_paragraph("End of doc.")

    text = extract_docx(_build_docx(build))
    assert "| Name | Age |" in text
    assert "| Alice | 30 |" in text
    assert "| Bob | 25 |" in text
    assert text.count("Alice") == 1


def test_paragraph_table_paragraph_stays_in_true_document_order() -> None:
    # python-docx's own .paragraphs and .tables are two separate flat
    # lists that do NOT preserve their relative order -- this is the
    # exact bug extraction_docx.py's module docstring documents finding
    # live before writing the real extraction logic.
    def build(d: Document) -> None:
        d.add_paragraph("Intro paragraph before the table.")
        table = d.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "cell"
        d.add_paragraph("Outro paragraph after the table.")

    text = extract_docx(_build_docx(build))
    assert text.index("Intro paragraph") < text.index("| cell |")
    assert text.index("| cell |") < text.index("Outro paragraph")


def test_empty_document_does_not_crash() -> None:
    def build(d: Document) -> None:
        pass

    content = _build_docx(build)
    assert isinstance(extract_docx(content), str)
