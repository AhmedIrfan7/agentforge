"""DOCX text/structure extraction (roadmap step 092) -- registered into
extraction.py:HANDLERS under "docx".

python-docx (MIT). Iterates document.element.body's raw XML children
directly, then wraps each one back into a Paragraph/Table object,
rather than using the library's own top-level .paragraphs/.tables lists
-- those are two separate flat lists that don't preserve the true
relative order between paragraphs and tables. Verified against a real
generated docx: a table placed between two paragraphs came back AFTER
both of them when read via .paragraphs then .tables, even though it
sits in the middle of the real document. Same table/text-ordering
problem extraction_pdf.py solves for PDF, for a different underlying
reason (there, table cells double as ordinary words; here, two
unrelated flat lists don't share a common order).

Headings map to markdown levels from the paragraph's real style name
("Title" -> #, "Heading N" -> N "#" characters, capped at 6, markdown's
own max) -- unlike PDF, docx encodes heading levels explicitly, so this
doesn't need a font-size heuristic.
"""

import io
import re

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

from extraction_tables import rows_to_markdown

_HEADING_LEVEL = re.compile(r"^Heading (\d+)$")


def _heading_prefix(style_name: str) -> str | None:
    if style_name == "Title":
        return "# "
    match = _HEADING_LEVEL.match(style_name)
    if match:
        level = min(int(match.group(1)), 6)
        return "#" * level + " "
    return None


def extract_docx(content: bytes) -> str:
    document = docx.Document(io.BytesIO(content))
    parts: list[str] = []
    for child in document.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            paragraph = Paragraph(child, document)
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else None
            prefix = _heading_prefix(style_name) if style_name is not None else None
            parts.append((prefix or "") + text)
        elif tag == "tbl":
            table = Table(child, document)
            markdown_table = rows_to_markdown(
                [[cell.text for cell in row.cells] for row in table.rows]
            )
            if markdown_table:
                parts.append(markdown_table)
    return "\n\n".join(parts)
